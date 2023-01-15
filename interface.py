
from pyChatGPT import ChatGPT
from os.path import join as opj
import os,time,datetime
import docx,csv
from docx.oxml.ns import qn
from docx.shared import Pt,RGBColor


import codecs

import warnings
warnings.filterwarnings("ignore")

def check_utf8(file_path):
    # 打开文件并读取前两行
    with open(file_path, 'rb') as f:
        lines = f.readlines()[:2]

    # 尝试使用 UTF-8 编码解码
    try:
        for line in lines:
            codecs.decode(line, 'utf-8')
        return True
    except:
        return False

def convert_utf8(file_path):
    # 打开文件并读取所有行
    with open(file_path, 'rb') as f:
        lines = f.readlines()

    # 使用 GBK 编码解码，然后使用 UTF-8 编码重新编码
    lines = [codecs.decode(line, 'gbk').encode('utf-8') for line in lines]

    # 打开新文件并使用 UTF-8 编码写入
    with open(file_path, 'wb') as f:
        f.writelines(lines)


baseoutput='./outputs'
titledir=opj(baseoutput,'任务单')

taskfile='/path/to/taskfile'

def readtasks():
    global taskfile
    list=os.listdir(titledir)
    list.sort(key=lambda fn: os.path.getmtime(opj(titledir,fn)) if not os.path.isdir(opj(titledir,fn)) else 0)
    d=datetime.datetime.fromtimestamp(os.path.getmtime(opj(titledir,list[-1])))
    print('正在读取最新任务文件：'+list[-1]+"，该文件最后保存时间："+d.strftime("%Y-%m-%d %H-%M-%S"))
    taskfile = opj(titledir, list[-1])
    if not check_utf8(taskfile):
        print('文件不是 UTF-8 编码，将进行转换')
        convert_utf8(taskfile)
        print('文件转换完成')
    with open(taskfile, 'r',encoding = 'utf-8') as f:
        reader=csv.reader(f)
        rows=[row for row in reader]
    return rows[1:]

#输入str类型的tid 读取csv找到并删除对应的这一条题目
def RMfromcsv(tid:str):
    if not os.path.exists(taskfile):
        return
    with open(taskfile, 'r', encoding='utf-8') as f:
        reader=csv.reader(f)
        rows=[row for row in reader]
    
    rest=[]
    for r in rows:
        if(r==[]):
            continue
        if r[0]==tid:
            print('移除了 这一个条目 from csv file: ', r)
        else:
            rest.append(r)

    with open(taskfile, 'w', encoding='utf-8', newline='') as f:
        writer=csv.writer(f)
        writer.writerows(rest)

from typing import List
def multidocs(tasklists:List[list])->bool:
    print('开始生产文档')
    ###########  init chatgpt ###############
    # read token
    with open('./token.txt','r') as f:
        cont=f.readlines()
    session_token=''
    for con in cont:
        session_token+=con.strip()
    # api1 = ChatGPT(session_token)  # auth with session token
    # api1 = ChatGPT(session_token, proxy='http://127.0.0.1:7890')  # specify proxy
    # api3 = ChatGPT(auth_type='google', email='example@gmail.com', password='password') # auth with google login
    api1 = ChatGPT(session_token)  # verbose mode (print debug messages)

    print('...')
    today=datetime.date.today()
    foldername = today.strftime('%y%m%d')
    outputdir=opj(baseoutput, foldername)
    if not os.path.exists(outputdir):
        os.makedirs(outputdir);
    idx=0
    while idx<len(tasklists):
        task=tasklists[idx]
        title=task[1] # format: [id,title]
        tid = task[0]
        if len(title)<5:
            print('标题太短不符合要求，废弃！')
            RMfromcsv(tid)
            idx+=1
            continue

        fdd = docx.Document()
        fdd.styles['Normal'].font.name = u'宋体'
        fdd.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        fdd.styles['Normal'].font.size = Pt(12)
        fdd.styles['Normal'].font.color.rgb = RGBColor(0,0,0)

        fdd.add_heading(title,0)

        zhengwen=''
        i=0
        try:
            resp = api1.send_message('请写一篇不少于800字的文章，但是内容形式和结构不要太单调，最好有一定的长度和深度，不要涉及政治。'
                '可以分小节论述，但是“小节一”、“小标题”之类的词语不要出现。'
                '不用说其他无关的话，大部分用中文，文章里不可以出现网址链接，如果写完了请在末尾加‘#123456’。'
                '文章的主题是: '+title)
        except Exception as e:
            print('碰到错误：',e)
            # time.sleep(300)
            api1 = ChatGPT(session_token)  # auth with session token
            continue
            
        while(resp['message'].find('#123456')==-1 and i<5):
            zhengwen+=resp['message'].strip()
            try:
                resp = api1.send_message('继续')
            except Exception as e:
                print('碰到错误：',e)
                # api1 = ChatGPT(session_token)  # auth with session token
                api1.reset_conversation()
                continue
            i+=1
        zhengwen+=resp['message'].replace('#123456',' ').strip()
        zhengwen.replace('\r','\n')
        fdd.add_paragraph(zhengwen)

        art=title+'.docx'
        RMfromcsv(tid)
        try:
            fdd.save(opj(outputdir,art))
        except Exception as e:
            print(e)
            print('这个出错了，下一个：',art)
        idx+=1
        api1.reset_conversation()  # reset the conversation
        time.sleep(3) #3 sec

    # close gpt
    # time.sleep(120) # wait 120 seconds 以避免网站的 1 hour limit
    api1.reset_conversation()  # reset the conversation
    api1.refresh_chat_page()
    return True



if __name__ =="__main__":
    res = False
    while True:
        try:
            tasks=readtasks()
            res = multidocs(tasks)
        except:
            time.sleep(5)
            if res is True:
                break

   

